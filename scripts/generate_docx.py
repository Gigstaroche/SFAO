from docx import Document
from docx.shared import Pt

SRC_MD = "../PROJECT_REPORT.md"
OUT_DOCX = "../PROJECT_REPORT.docx"

def add_preformatted(paragraph, text):
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(9)


def main():
    doc = Document()
    try:
        with open(SRC_MD, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Source markdown not found: {SRC_MD}")
        return

    in_code = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            in_code = not in_code
            if not in_code:
                # flush code block
                p = doc.add_paragraph()
                for cl in code_lines:
                    add_preformatted(p, cl + '\n')
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == '---':
            doc.add_paragraph('')
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    # Add explicit placeholders for images and screenshots
    doc.add_heading('Screenshots & Image Placeholders', level=2)
    placeholders = [
        ('Dashboard landing', 'frontend/assets/dashboard.png'),
        ('Sample sentiment chart (doughnut)', 'frontend/assets/sentiment_doughnut.png'),
        ('Database Studio table view', 'frontend/assets/database_table.png'),
        ('Architecture diagram', 'frontend/assets/architecture.png'),
    ]
    for title, path in placeholders:
        p = doc.add_paragraph()
        p.add_run('PLACEHOLDER: ').bold = True
        p.add_run(f"Insert image for: {title}\nPath suggestion: {path}")

    doc.add_heading('Image Guidelines', level=2)
    doc.add_paragraph('- Recommended image formats: PNG or JPG')
    doc.add_paragraph('- Place images in the suggested paths, then reopen the document to replace placeholders or send them to the reviewer.')

    doc.add_heading('Appendix: Files Referenced', level=2)
    files = [
        'README.md',
        'DATABASE_INTEGRATION.md',
        'PORTAL_PHASE_PLAN.md',
        'prisma/schema.prisma',
        'backend/brain.py',
        'backend/database.py',
        'backend/main.py',
    ]
    for f in files:
        doc.add_paragraph(f)

    doc.save(OUT_DOCX)
    print(f"Saved {OUT_DOCX}")

if __name__ == '__main__':
    main()
