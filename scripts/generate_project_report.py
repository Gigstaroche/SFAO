from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Image


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "PROJECT_REPORT.docx"
PDF_PATH = ROOT / "PROJECT_REPORT.pdf"
SCREENSHOT_PATH = ROOT / "dashboard-real-viewport.png"


def set_run_font(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_docx_placeholder(document, title, lines):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FF")
    set_cell_border(
        cell,
        top={"val": "single", "sz": 12, "color": "B7C4E8"},
        bottom={"val": "single", "sz": 12, "color": "B7C4E8"},
        left={"val": "single", "sz": 12, "color": "B7C4E8"},
        right={"val": "single", "sz": 12, "color": "B7C4E8"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True, color=(42, 58, 115))
    for line in lines:
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = para.add_run(line)
        set_run_font(rr, size=10, color=(60, 60, 60))
    document.add_paragraph()


def build_docx():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SFAO - Smart Feedback Analyzer for Organization")
    set_run_font(r, size=20, bold=True, color=(28, 44, 90))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dubai Future Solutions - Prototypes for Humanity Application Report")
    set_run_font(r, size=13, bold=True, color=(76, 95, 145))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Unified offline intelligence for privacy-first organizational feedback analysis")
    set_run_font(r, size=11, color=(75, 75, 75))

    doc.add_paragraph("Executive Summary")
    doc.add_paragraph(
        "SFAO is a privacy-first feedback intelligence platform built to help organizations understand internal and external feedback without sending sensitive data to cloud services. "
        "It combines local NLP processing, multi-source ingestion, actionable dashboards, and operational monitoring in one offline system. "
        "For the Dubai Future Solutions - Prototypes for Humanity summit, SFAO presents a practical prototype that addresses a real organizational problem: feedback is abundant, but decision-makers need a secure and efficient way to convert it into action."
    )

    for title, bullets in [
        (
            "Problem and Opportunity",
            [
                "Organizations collect feedback from surveys, emails, service desks, and social channels, but the information is fragmented and often reviewed too late.",
                "Cloud AI tools can be expensive, difficult to govern, and unsuitable for sensitive internal content.",
                "SFAO solves this by running analysis locally, preserving privacy while still producing management-ready insights.",
            ],
        ),
        (
            "Why This Fits Prototypes for Humanity",
            [
                "The project addresses a real-world productivity and governance problem rather than a purely theoretical concept.",
                "It is already implemented as a working prototype with a dashboard, survey flow, database layer, and monitoring tools.",
                "The concept is scalable, socially useful, and designed for deployment in schools, SMEs, nonprofits, and enterprise teams that need low-cost analytics.",
            ],
        ),
        (
            "Core Capabilities",
            [
                "Local sentiment analysis and urgency detection.",
                "Feedback aggregation from surveys and simulated social sources.",
                "Executive dashboard with charts, filters, and action tracking.",
                "Admin monitoring and audit logs for operational visibility.",
                "SQLite-based offline storage for secure and portable deployment.",
            ],
        ),
        (
            "Technology Stack",
            [
                "Backend: FastAPI.",
                "Database: SQLite.",
                "AI / NLP: VADER sentiment analysis.",
                "Frontend: HTML, CSS, JavaScript, Chart.js.",
                "Automation: Python scripts for seeding and feedback simulation.",
            ],
        ),
    ]:
        doc.add_heading(title, level=1)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("System Architecture", level=1)
    doc.add_paragraph(
        "The architecture is deliberately lightweight so it can run on a single local machine or inside a small institutional deployment without external infrastructure."
    )
    arch = doc.add_table(rows=4, cols=2)
    arch.style = "Table Grid"
    arch_rows = [
        ("Frontend", "Executive dashboard, survey portal, and admin interfaces built with HTML, CSS, and JavaScript."),
        ("Backend", "FastAPI routes, authentication, sentiment processing, and monitoring endpoints."),
        ("Data Layer", "SQLite storage for feedback records, users, actions, and monitoring logs."),
        ("AI Layer", "Local VADER-based NLP pipeline for sentiment scoring and urgency detection."),
    ]
    for row, (left, right) in zip(arch.rows, arch_rows):
        row.cells[0].text = left
        row.cells[1].text = right

    doc.add_heading("Impact and Value", level=1)
    for bullet in [
        "Improves privacy by keeping data processing local.",
        "Reduces recurring AI service costs by avoiding cloud inference APIs.",
        "Helps executives identify issues early through automated trend analysis.",
        "Creates a closed feedback loop from submission to resolution.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Future Roadmap", level=1)
    for bullet in [
        "Add stronger text categorization for more accurate routing.",
        "Expand anonymization and privacy controls before storage.",
        "Add lightweight local language-model summaries for management reports.",
        "Integrate additional institutional workflows such as case assignment and SLA tracking.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Recommended Presentation Structure", level=1)
    for bullet in [
        "Start with the problem: feedback is fragmented, slow to analyze, and sensitive.",
        "Show the prototype working live: dashboard, survey intake, and monitoring.",
        "End with impact: privacy, affordability, and clear organizational value.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Image Placeholders", level=1)
    doc.add_paragraph("Insert your own screenshots before final submission. Use the placeholder names below in the PDF.")
    add_docx_placeholder(
        doc,
        "Figure 1 - Executive Dashboard Screenshot",
        [
            "Insert a screenshot of the signed-in dashboard view.",
            "Show the charts, live feed, and workboard area.",
        ],
    )
    add_docx_placeholder(
        doc,
        "Figure 2 - Employee Survey Portal Screenshot",
        [
            "Insert a screenshot of the feedback form page.",
            "Show the form fields and submission controls.",
        ],
    )
    add_docx_placeholder(
        doc,
        "Figure 3 - Admin Monitoring Panel Screenshot",
        [
            "Insert a screenshot of the monitoring or audit view.",
            "Show the event list and export controls.",
        ],
    )
    add_docx_placeholder(
        doc,
        "Figure 4 - Database Studio Screenshot",
        [
            "Insert a screenshot of the database studio or query page.",
            "Show tables, records, or query output.",
        ],
    )

    doc.add_heading("Application Notes", level=1)
    doc.add_paragraph(
        "For the Dubai Future Solutions application, keep the story focused on a live working prototype, a clear problem, and the potential for real-world deployment. "
        "The strongest angle is that SFAO demonstrates how local AI can make organizational feedback more actionable without compromising privacy."
    )

    doc.save(DOCX_PATH)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#1c2c5a")))
    styles.add(ParagraphStyle(name="SubtitleCenter", parent=styles["Normal"], alignment=TA_CENTER, fontName="Helvetica", fontSize=11, leading=14, textColor=colors.HexColor("#4c5f91")))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#1c2c5a"), spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=5))
    styles.add(ParagraphStyle(name="BulletBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, leftIndent=12, bulletIndent=0, spaceAfter=3))
    styles.add(ParagraphStyle(name="FigureTitle", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=11, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#2a3a73")))
    styles.add(ParagraphStyle(name="FigureBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12, alignment=TA_CENTER, textColor=colors.HexColor("#3c3c3c")))

    story = []

    def bullet(text):
        story.append(Paragraph(f"• {text}", styles["BulletBody"]))

    story.append(Paragraph("SFAO - Smart Feedback Analyzer for Organization", styles["TitleCenter"]))
    story.append(Spacer(1, 2))
    story.append(Paragraph("Dubai Future Solutions - Prototypes for Humanity Application Report", styles["SubtitleCenter"]))
    story.append(Paragraph("Unified offline intelligence for privacy-first organizational feedback analysis", styles["SubtitleCenter"]))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Executive Summary", styles["Section"]))
    story.append(Paragraph(
        "SFAO is a privacy-first feedback intelligence platform built to help organizations understand internal and external feedback without sending sensitive data to cloud services. "
        "It combines local NLP processing, multi-source ingestion, actionable dashboards, and operational monitoring in one offline system. "
        "For the Dubai Future Solutions - Prototypes for Humanity summit, SFAO presents a practical prototype that addresses a real organizational problem: feedback is abundant, but decision-makers need a secure and efficient way to convert it into action.",
        styles["Body"]
    ))

    for title, bullets in [
        (
            "Problem and Opportunity",
            [
                "Organizations collect feedback from surveys, emails, service desks, and social channels, but the information is fragmented and often reviewed too late.",
                "Cloud AI tools can be expensive, difficult to govern, and unsuitable for sensitive internal content.",
                "SFAO solves this by running analysis locally, preserving privacy while still producing management-ready insights.",
            ],
        ),
        (
            "Why This Fits Prototypes for Humanity",
            [
                "The project addresses a real-world productivity and governance problem rather than a purely theoretical concept.",
                "It is already implemented as a working prototype with a dashboard, survey flow, database layer, and monitoring tools.",
                "The concept is scalable, socially useful, and designed for deployment in schools, SMEs, nonprofits, and enterprise teams that need low-cost analytics.",
            ],
        ),
        (
            "Core Capabilities",
            [
                "Local sentiment analysis and urgency detection.",
                "Feedback aggregation from surveys and simulated social sources.",
                "Executive dashboard with charts, filters, and action tracking.",
                "Admin monitoring and audit logs for operational visibility.",
                "SQLite-based offline storage for secure and portable deployment.",
            ],
        ),
        (
            "Technology Stack",
            [
                "Backend: FastAPI.",
                "Database: SQLite.",
                "AI / NLP: VADER sentiment analysis.",
                "Frontend: HTML, CSS, JavaScript, Chart.js.",
                "Automation: Python scripts for seeding and feedback simulation.",
            ],
        ),
    ]:
        story.append(Paragraph(title, styles["Section"]))
        for item in bullets:
            bullet(item)

    story.append(Paragraph("System Architecture", styles["Section"]))
    story.append(Paragraph(
        "The architecture is deliberately lightweight so it can run on a single local machine or inside a small institutional deployment without external infrastructure.",
        styles["Body"]
    ))
    arch = Table(
        [
            ["Frontend", "Executive dashboard, survey portal, and admin interfaces built with HTML, CSS, and JavaScript."],
            ["Backend", "FastAPI routes, authentication, sentiment processing, and monitoring endpoints."],
            ["Data Layer", "SQLite storage for feedback records, users, actions, and monitoring logs."],
            ["AI Layer", "Local VADER-based NLP pipeline for sentiment scoring and urgency detection."],
        ],
        colWidths=[34 * mm, 136 * mm],
    )
    arch.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF3FF")),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#9FB0D9")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9D4EE")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("LEADING", (0, 0), (-1, -1), 12),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(arch)
    story.append(Spacer(1, 6))

    story.append(Paragraph("Impact and Value", styles["Section"]))
    for item in [
        "Improves privacy by keeping data processing local.",
        "Reduces recurring AI service costs by avoiding cloud inference APIs.",
        "Helps executives identify issues early through automated trend analysis.",
        "Creates a closed feedback loop from submission to resolution.",
    ]:
        bullet(item)

    story.append(Paragraph("Future Roadmap", styles["Section"]))
    for item in [
        "Add stronger text categorization for more accurate routing.",
        "Expand anonymization and privacy controls before storage.",
        "Add lightweight local language-model summaries for management reports.",
        "Integrate additional institutional workflows such as case assignment and SLA tracking.",
    ]:
        bullet(item)

    story.append(Paragraph("Recommended Presentation Structure", styles["Section"]))
    for item in [
        "Start with the problem: feedback is fragmented, slow to analyze, and sensitive.",
        "Show the prototype working live: dashboard, survey intake, and monitoring.",
        "End with impact: privacy, affordability, and clear organizational value.",
    ]:
        bullet(item)

    story.append(Paragraph("Image Placeholders", styles["Section"]))
    story.append(Paragraph("Insert your own screenshots before final submission. The placeholder names below are intentional and can be replaced with your captured images.", styles["Body"]))

    placeholders = [
        ("Figure 1 - Executive Dashboard Screenshot", ["Insert a screenshot of the signed-in dashboard view.", "Show the charts, live feed, and workboard area."] ),
        ("Figure 2 - Employee Survey Portal Screenshot", ["Insert a screenshot of the feedback form page.", "Show the form fields and submission controls."] ),
        ("Figure 3 - Admin Monitoring Panel Screenshot", ["Insert a screenshot of the monitoring or audit view.", "Show the event list and export controls."] ),
        ("Figure 4 - Database Studio Screenshot", ["Insert a screenshot of the database studio or query page.", "Show tables, records, or query output."] ),
    ]

    for title, lines in placeholders:
        story.append(Paragraph(title, styles["FigureTitle"]))
        data = [[Paragraph(line, styles["FigureBody"])] for line in lines]
        table = Table(data, colWidths=[170 * mm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F7FF")),
            ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#B7C4E8")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D6DDF2")),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(table)
        story.append(Spacer(1, 8))

    story.append(Paragraph("Submission Guidance", styles["Section"]))
    story.append(Paragraph(
        "For the Dubai Future Solutions application, keep the narrative focused on a live working prototype, a clear problem, and the potential for real-world deployment. "
        "The strongest angle is that SFAO demonstrates how local AI can make organizational feedback more actionable without compromising privacy.",
        styles["Body"]
    ))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    doc.build(story)


def main():
    build_docx()
    build_pdf()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()